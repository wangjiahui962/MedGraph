from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK


OUT = "/Users/jiahui/Documents/project/MedGraph/docs/MedGraph_智能体驱动的知识图谱自动构建系统_设计方案.docx"

NAVY = "17365D"
BLUE = "2F75B5"
PALE = "EAF2F8"
LIGHT = "F3F6F9"
GRAY = "5B6573"
WHITE = "FFFFFF"
GREEN = "2E7D32"
ORANGE = "C76B00"
RED = "A61B1B"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, name="Hiragino Sans GB", size=None, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, size=9, color=GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    tail = paragraph.add_run(" 页")
    set_font(tail, size=9, color=GRAY)


def setup_styles(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.78)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.88)
    sec.right_margin = Inches(0.88)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Hiragino Sans GB"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    specs = {
        "Title": (28, NAVY, 0, 8),
        "Subtitle": (13, GRAY, 0, 10),
        "Heading 1": (16, NAVY, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (11.5, NAVY, 9, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Hiragino Sans GB"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        s = doc.styles[style_name]
        s.font.name = "Hiragino Sans GB"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
        s.font.size = Pt(10.5)
        s.paragraph_format.left_indent = Inches(0.5)
        s.paragraph_format.first_line_indent = Inches(-0.25)
        s.paragraph_format.space_after = Pt(4)
        s.paragraph_format.line_spacing = 1.15

    if "Callout" not in [s.name for s in doc.styles]:
        s = doc.styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = normal
        s.font.name = "Hiragino Sans GB"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
        s.font.size = Pt(10.5)
        s.font.color.rgb = RGBColor.from_string(NAVY)
        s.paragraph_format.left_indent = Inches(0.18)
        s.paragraph_format.right_indent = Inches(0.12)
        s.paragraph_format.space_before = Pt(6)
        s.paragraph_format.space_after = Pt(8)


def add_header_footer(doc):
    sec = doc.sections[0]
    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("MEDGRAPH  /  系统设计方案")
    set_font(r, size=8.5, bold=True, color=GRAY)
    p_pr = header._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "D9E2F3")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    add_page_number(sec.footer.paragraphs[0])


def add_para(doc, text="", bold_prefix=None, style=None, color=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2, color=color)
    else:
        r = p.add_run(text)
        set_font(r, color=color)
    return p


def add_bullets(doc, items):
    for item in items:
        add_para(doc, item, style="List Bullet")


def add_numbers(doc, items):
    for item in items:
        add_para(doc, item, style="List Number")


def add_callout(doc, label, text, fill=PALE):
    p = doc.add_paragraph(style="Callout")
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    r = p.add_run(label + "  ")
    set_font(r, bold=True, color=NAVY)
    r = p.add_run(text)
    set_font(r, color=NAVY)


def add_table(doc, headers, rows, widths, header_fill=BLUE, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_font(r, size=font_size, bold=True, color=WHITE if header_fill != LIGHT else NAVY)
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            if ri % 2 == 1:
                set_cell_shading(cells[idx], "F8FAFC")
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(text))
            set_font(r, size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()
    setup_styles(doc)
    add_header_footer(doc)

    # Cover
    for _ in range(3):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(kicker.add_run("题目三 · 总体设计方案"), size=11, bold=True, color=BLUE)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(title.add_run("智能体驱动的知识图谱\n自动构建系统"), size=28, bold=True, color=NAVY)
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(subtitle.add_run("MedGraph：面向医疗文本的数据采集、知识抽取、图谱构建与可视化平台"), size=13, color=GRAY)
    doc.add_paragraph()
    add_callout(doc, "方案定位", "以多智能体工作流贯通数据采集、清洗、实体关系抽取、质量审查、图谱入库与可视化展示；按可验收口径设计，目标覆盖 120 个类别、3,600 条文本，形成可追溯的实体、属性与关系三元组。")
    doc.add_paragraph()
    meta = add_table(doc, ["项目", "内容"], [
        ("文档性质", "立项与技术设计方案（暂不包含实现代码）"),
        ("建议版本", "V1.0"),
        ("适用阶段", "需求评审、技术选型、任务分工、答辩准备"),
        ("建议周期", "8 周，可按课程节点压缩为 6 周"),
    ], [1900, 7460], header_fill=NAVY, font_size=9.5)
    add_para(doc, "编制日期：2026 年 9 月", color=GRAY).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_break(doc)

    doc.add_heading("1. 项目概述", level=1)
    add_callout(doc, "一句话方案", "构建一个“可配置、可追踪、可复核”的多智能体流水线：自动采集医疗文本，经规则 + 领域模型 + 大语言模型协同抽取知识，写入图数据库，并通过 Web 界面提供浏览、查询和关系展示。")
    doc.add_heading("1.1 建设目标", level=2)
    add_bullets(doc, [
        "规模目标：覆盖不少于 120 个医疗类别，采集并保留不少于 3,600 条有效文本；最终验收底线仍按题目要求的 100 类、3,000 条计算。",
        "能力目标：自动识别实体、实体属性和实体关系，生成结构化三元组并完成去重、校验、溯源和入库。",
        "系统目标：实现智能体调度、任务状态跟踪、失败重试、人工复核入口和增量更新机制。",
        "展示目标：提供图谱全局浏览、实体查询、邻居展开、关系路径展示、来源证据查看和基础统计面板。",
        "交付目标：形成完整源代码、知识图谱数据文件、运行说明、评测报告和项目展示 PPT。",
    ])
    doc.add_heading("1.2 设计原则", level=2)
    add_table(doc, ["原则", "设计含义"], [
        ("可验收", "每项题目要求均映射到数据指标、接口、页面或交付文件。"),
        ("可追溯", "实体和关系保存来源 URL、文本片段、采集时间、抽取模型与置信度。"),
        ("可扩展", "类别、采集器、模型、提示词、图谱模式和存储后端均可配置。"),
        ("可复核", "低置信度、冲突知识和异常样本进入人工审核队列。"),
        ("合规优先", "只采集公开授权或允许使用的数据，遵守 robots、许可协议和隐私要求。"),
    ], [1800, 7560])

    doc.add_heading("2. 题目要求与验收口径", level=1)
    add_table(doc, ["题目要求", "方案响应", "建议验收证据"], [
        ("自动完成四个阶段", "采集、抽取、构建、展示由智能体工作流串联", "任务日志、运行录像、流程状态页"),
        ("不少于 100 个类别", "规划 120 个疾病/主题类别", "category.csv、类别统计页"),
        ("不少于 3,000 条文本", "规划 3,600 条清洗后有效文档", "documents.jsonl、哈希去重报告"),
        ("识别实体、属性、关系", "规则 + NER/RE 模型 + LLM 结构化抽取", "标注样例、评测集、抽取结果"),
        ("形成三元组并存储", "RDF/CSV 导出 + Neo4j 图数据库", "triples.csv、Neo4j 数据库截图"),
        ("图谱浏览与查询", "Web 可视化、实体详情、邻居与路径展示", "可运行页面和功能测试"),
        ("完整交付", "代码、数据、文档、PPT、启动脚本", "交付目录与 README"),
    ], [2100, 4050, 3210], font_size=8.7)
    add_callout(doc, "统计口径约定", "“类别数”以 category_id 去重；“文本数”以清洗后文档 document_id 去重；“实体数”以规范化 entity_id 去重；“三元组数”以 (subject_id, predicate, object_id/value) 去重。四项必须分别统计。", fill="FFF3CD")

    doc.add_heading("3. 业务范围与数据方案", level=1)
    doc.add_heading("3.1 领域选择", level=2)
    add_para(doc, "建议选择“常见疾病与临床知识”作为 R 大类文本的业务主题，将类别层定义为疾病类别或临床专题。医疗领域实体类型稳定、关系语义清晰，适合展示知识图谱价值，也与项目名 MedGraph 一致。若题目中的“R 大类”有特定赛题含义，只需替换领域配置与本体，不改变平台架构。")
    doc.add_heading("3.2 类别与数据规模规划", level=2)
    add_table(doc, ["一级主题", "示例类别", "计划类别数", "每类文本", "计划文本量"], [
        ("内科系统", "呼吸、消化、心血管、内分泌等", "32", "30", "960"),
        ("外科与急症", "普外、骨科、神外、创伤等", "20", "30", "600"),
        ("妇儿与生殖", "妇科、产科、儿科、新生儿等", "16", "30", "480"),
        ("专科疾病", "眼科、耳鼻喉、口腔、皮肤等", "20", "30", "600"),
        ("感染与免疫", "传染病、风湿免疫、过敏等", "16", "30", "480"),
        ("肿瘤与慢病管理", "常见肿瘤、康复、营养等", "16", "30", "480"),
        ("合计", "120 个可配置类别", "120", "平均 30", "3,600"),
    ], [1800, 3300, 1300, 1300, 1660], font_size=8.6)
    doc.add_heading("3.3 数据源策略", level=2)
    add_bullets(doc, [
        "首选公开、稳定、可引用的数据源：政府/医院科普页面、开放医学百科、开放论文摘要、公开药品或疾病说明数据集。",
        "每个来源配置独立采集器，限制抓取频率，并记录许可证、robots 策略、URL、发布时间和采集时间。",
        "对动态页面使用浏览器渲染采集；对静态页面优先 HTTP 抓取；对开放数据集优先直接下载并校验哈希。",
        "不采集患者隐私、付费墙内容或禁止自动抓取的站点；演示数据保留证据片段但不展示敏感信息。",
    ])
    doc.add_heading("3.4 文档数据结构", level=2)
    add_table(doc, ["字段", "说明", "示例"], [
        ("document_id", "清洗后文档唯一标识", "doc_sha256_..."),
        ("category_id", "类别标识，可多标签", "respiratory_asthma"),
        ("title / content", "标题与正文", "支气管哮喘的症状……"),
        ("source_url", "来源地址", "https://..."),
        ("license", "使用许可或来源声明", "open / public-info"),
        ("collected_at", "采集时间", "2026-09-01T..."),
        ("content_hash", "正文哈希，用于去重", "sha256"),
        ("quality_score", "清洗后质量分", "0.92"),
    ], [1900, 3800, 3660], font_size=8.7)

    add_page_break(doc)
    doc.add_heading("4. 系统总体架构", level=1)
    add_para(doc, "系统采用前后端分离、任务队列驱动和双存储架构。关系数据库保存任务、文档、审核与统计元数据；图数据库保存实体和关系；对象/文件存储保留原文、抽取结果和可导出的数据文件。")
    doc.add_heading("4.1 分层架构", level=2)
    add_table(doc, ["层级", "核心组件", "职责"], [
        ("展示层", "Vue 3 / React、ECharts 或 Cytoscape.js", "图谱浏览、查询、详情、路径、统计、任务状态"),
        ("服务层", "FastAPI", "统一 REST API、鉴权、查询编排、文件导出"),
        ("智能体编排层", "LangGraph/自研状态机 + Celery/RQ", "多智能体协作、状态持久化、重试和断点续跑"),
        ("处理层", "采集器、清洗器、NER、RE、LLM、实体链接", "数据获取与知识抽取"),
        ("存储层", "PostgreSQL/SQLite、Neo4j、JSONL/CSV", "元数据、图数据、原始与交付数据"),
        ("观测层", "结构化日志、指标、审计记录", "追踪任务耗时、错误、模型版本和质量"),
    ], [1500, 3400, 4460], font_size=8.8)
    doc.add_heading("4.2 端到端数据流", level=2)
    add_numbers(doc, [
        "任务规划：读取类别清单、来源配置和数量目标，生成分批采集任务。",
        "采集与清洗：抓取文本，抽取正文，进行语言检测、去噪、去重和质量评分。",
        "候选知识抽取：规则和领域模型识别实体，关系模型/LLM 输出符合 JSON Schema 的候选事实。",
        "规范化与融合：执行别名归一、实体链接、单位标准化、重复关系合并和冲突检测。",
        "质量审查：验证本体约束、证据跨度与置信度；低置信度结果转人工审核。",
        "图谱入库：通过幂等 Upsert 写入 Neo4j，同时导出 nodes.csv、edges.csv 和 triples.csv。",
        "查询与展示：API 将图查询结果转换为前端节点边格式，支持按需展开以控制渲染规模。",
    ])
    doc.add_heading("4.3 推荐技术栈", level=2)
    add_table(doc, ["模块", "推荐选型", "备选/说明"], [
        ("后端", "Python 3.11 + FastAPI", "生态成熟，便于调用 NLP/LLM"),
        ("智能体", "LangGraph 或轻量状态机", "课程项目建议显式状态机，便于解释"),
        ("异步任务", "Celery + Redis", "小规模演示可用 FastAPI BackgroundTasks"),
        ("NLP", "HanLP / Transformers", "结合词典和正则规则"),
        ("LLM", "可替换的模型适配器", "云端 API 或本地 Qwen 类模型"),
        ("图数据库", "Neo4j Community", "演示友好，Cypher 查询直观"),
        ("元数据", "PostgreSQL", "开发期可先用 SQLite"),
        ("前端", "Vue 3 + TypeScript + Cytoscape.js", "ECharts Graph 亦可"),
        ("部署", "Docker Compose", "一条命令启动服务依赖"),
    ], [1600, 3200, 4560], font_size=8.8)

    doc.add_heading("5. 多智能体设计", level=1)
    add_table(doc, ["智能体", "输入", "主要动作", "输出/终止条件"], [
        ("Orchestrator", "目标类别与数量", "拆解任务、调度、重试、汇总", "所有批次达标或触发人工介入"),
        ("Collector", "来源与采集策略", "请求、解析、限速、分页", "原始文档与采集日志"),
        ("Cleaner", "原始 HTML/文本", "正文提取、去噪、去重、评分", "合格标准文档"),
        ("Extractor", "标准文档、本体", "NER、属性与关系抽取", "候选实体和三元组"),
        ("Linker", "候选实体、词典", "别名归一、实体链接、消歧", "规范实体 ID"),
        ("Reviewer", "候选知识与证据", "规则校验、冲突检测、置信度融合", "通过、驳回或人工复核"),
        ("GraphWriter", "审核通过的知识", "幂等写入、版本与溯源", "节点边及导出文件"),
        ("Reporter", "全流程指标", "统计规模、质量、失败原因", "验收报告与看板数据"),
    ], [1450, 1950, 3100, 2860], font_size=8.2)
    add_callout(doc, "智能体边界", "智能体不是简单地“调用一次大模型”。每个智能体具有明确的目标、输入输出 Schema、可用工具、状态、最大重试次数与停止条件；编排器只在状态转换中传递结构化数据。")
    doc.add_heading("5.1 状态机与容错", level=2)
    add_bullets(doc, [
        "主状态：PENDING → COLLECTED → CLEANED → EXTRACTED → REVIEWED → STORED → COMPLETED。",
        "异常状态：RETRYABLE、REJECTED、NEEDS_REVIEW、FAILED；每一步保存检查点，支持断点续跑。",
        "幂等策略：document_id、entity_id、triple_key 建唯一约束；重复执行不产生重复节点或边。",
        "降级策略：LLM 不可用时保留规则与本地模型结果；Neo4j 不可用时先写入待导入文件。",
    ])
    doc.add_heading("5.2 提示词与结构化输出", level=2)
    add_para(doc, "LLM 抽取必须通过 JSON Schema 限制字段，禁止直接生成 Cypher。输出至少包含 subject、predicate、object、subject_type、object_type、evidence、confidence；随后由 Reviewer 校验实体类型、关系域/值域、证据是否出现在原文及置信度阈值。")

    doc.add_heading("6. 知识图谱模式设计", level=1)
    doc.add_heading("6.1 核心实体类型", level=2)
    add_table(doc, ["实体类型", "关键属性", "示例"], [
        ("Disease 疾病", "name、aliases、definition、icd_code", "支气管哮喘"),
        ("Symptom 症状", "name、body_site、severity", "喘息"),
        ("Drug 药物", "generic_name、dosage_form", "布地奈德"),
        ("Treatment 治疗", "name、type", "吸入治疗"),
        ("Examination 检查", "name、method", "肺功能检查"),
        ("Department 科室", "name", "呼吸内科"),
        ("Population 人群", "name、age_range", "儿童"),
        ("RiskFactor 危险因素", "name、category", "过敏原暴露"),
        ("Complication 并发症", "name", "呼吸衰竭"),
    ], [2200, 4200, 2960], font_size=8.7)
    doc.add_heading("6.2 核心关系类型", level=2)
    add_table(doc, ["关系", "主语 → 宾语", "示例"], [
        ("HAS_SYMPTOM", "Disease → Symptom", "哮喘 → 喘息"),
        ("TREATED_BY", "Disease → Drug/Treatment", "哮喘 → 布地奈德"),
        ("DIAGNOSED_BY", "Disease → Examination", "哮喘 → 肺功能检查"),
        ("BELONGS_TO", "Disease → Department", "哮喘 → 呼吸内科"),
        ("HIGH_RISK_FOR", "Population → Disease", "过敏体质人群 → 哮喘"),
        ("HAS_RISK_FACTOR", "Disease → RiskFactor", "哮喘 → 过敏原暴露"),
        ("MAY_CAUSE", "Disease → Complication", "重症哮喘 → 呼吸衰竭"),
        ("CO_OCCURS_WITH", "Disease ↔ Disease", "哮喘 ↔ 过敏性鼻炎"),
    ], [2200, 3100, 4060], font_size=8.7)
    doc.add_heading("6.3 三元组与溯源模型", level=2)
    add_para(doc, "示例三元组：(支气管哮喘, HAS_SYMPTOM, 喘息)。关系边同时保存 confidence、source_document_id、evidence_text、extractor、model_version、created_at。实体属性采用键值方式存储，但在导出时也可转换为 (实体, 属性名, 属性值) 三元组，以满足题目对属性抽取的要求。")
    add_callout(doc, "版本策略", "知识不是覆盖式更新。对同一事实保留来源集合与首次/最近发现时间；冲突事实进入审查队列，最终状态标记为 accepted、disputed 或 deprecated。")

    add_page_break(doc)
    doc.add_heading("7. 信息抽取与知识融合方案", level=1)
    doc.add_heading("7.1 混合抽取路线", level=2)
    add_table(doc, ["阶段", "方法", "作用"], [
        ("候选识别", "医学词典、正则、Aho–Corasick", "高精度识别药名、疾病名、数值和单位"),
        ("实体识别", "中文预训练模型微调或现成 NER", "补充上下文相关实体和新词"),
        ("关系抽取", "规则模板 + 分类模型", "处理高频、稳定关系"),
        ("复杂抽取", "LLM few-shot + JSON Schema", "处理跨句关系、属性与长尾表达"),
        ("实体链接", "别名词典 + 字符/向量相似度", "归并同义词并连接规范实体"),
        ("质量融合", "规则投票 + 模型置信度 + 来源权重", "计算最终可信度并识别冲突"),
    ], [1700, 3500, 4160], font_size=8.8)
    doc.add_heading("7.2 质量控制规则", level=2)
    add_bullets(doc, [
        "Schema 校验：实体类型、关系名称、属性类型、关系域和值域必须符合本体。",
        "证据校验：evidence_text 必须能在原文定位，并保存起止字符位置。",
        "一致性校验：同一实体不得同时拥有明显互斥属性；数值单位统一后再比较。",
        "置信度策略：≥0.85 自动通过；0.60–0.85 进入抽样或人工复核；<0.60 默认不入库。",
        "抽样复核：每个类别至少复核固定数量样本，避免总体指标掩盖长尾类别质量问题。",
    ])
    doc.add_heading("7.3 评测集设计", level=2)
    add_para(doc, "从 120 个类别中分层抽取不少于 30 个类别，每类人工标注 5–10 条文本，形成 200–300 条金标准评测集。标注实体边界、实体类型、实体规范名和关系三元组，双人标注并解决分歧。")
    add_table(doc, ["指标", "计算/目标", "用途"], [
        ("实体 Precision / Recall / F1", "严格边界匹配；F1 建议 ≥ 0.80", "衡量实体识别质量"),
        ("关系 Precision / Recall / F1", "实体与关系均正确；F1 建议 ≥ 0.70", "衡量三元组质量"),
        ("实体链接准确率", "正确规范实体数 / 总链接数；≥ 0.85", "衡量去重与消歧"),
        ("证据可定位率", "可回溯三元组 / 总三元组；≥ 0.95", "保障可解释性"),
        ("重复率", "重复文档或三元组占比；≤ 0.02", "保障数据有效规模"),
        ("流程成功率", "成功批次 / 总批次；≥ 0.95", "衡量系统稳定性"),
    ], [2600, 4100, 2660], font_size=8.6)

    doc.add_heading("8. 图数据库与接口设计", level=1)
    doc.add_heading("8.1 Neo4j 存储", level=2)
    add_bullets(doc, [
        "节点以 entity_id 建唯一约束，name 和 aliases 建全文索引；类别与文档节点用于追溯。",
        "关系使用规范谓词名称，边属性保存置信度、证据、来源、抽取器和时间戳。",
        "写入采用批量 UNWIND + MERGE，避免逐条请求；所有写入操作支持重复执行。",
        "大图浏览默认限制深度、节点数和超时时间，防止一次查询加载整个图谱。",
    ])
    doc.add_heading("8.2 核心 API", level=2)
    add_table(doc, ["方法与路径", "功能", "关键参数/返回"], [
        ("POST /api/jobs", "创建构建任务", "category_ids、target_count、sources"),
        ("GET /api/jobs/{id}", "查看任务状态", "阶段、进度、错误、统计"),
        ("GET /api/entities/search", "实体搜索", "q、type、limit"),
        ("GET /api/entities/{id}", "实体详情", "属性、来源、相邻关系"),
        ("GET /api/graph/neighbors", "邻居展开", "entity_id、depth、relation_types"),
        ("GET /api/graph/path", "关系路径", "source_id、target_id、max_depth"),
        ("GET /api/stats", "规模与质量统计", "类别、文本、实体、边、指标"),
        ("GET /api/export", "导出数据", "format=csv/jsonl/rdf"),
    ], [2600, 2800, 3960], font_size=8.7)

    doc.add_heading("9. 可视化界面设计", level=1)
    add_table(doc, ["页面", "核心功能", "验收点"], [
        ("总览仪表盘", "类别、文本、实体、关系、任务进度和质量指标", "数字与数据文件统计一致"),
        ("图谱浏览", "节点拖拽、缩放、筛选、按需展开、布局切换", "可浏览且不卡死"),
        ("实体查询", "名称/别名模糊搜索、类型过滤、结果排序", "可查到实体并定位节点"),
        ("实体详情", "属性、邻居、来源证据、置信度", "信息完整、可追溯"),
        ("关系展示", "直接关系、按类型过滤、最短路径", "关系方向和类型清晰"),
        ("任务中心", "启动、暂停、状态、失败原因、重试", "体现智能体自动化过程"),
        ("人工审核", "候选三元组通过、驳回、修改", "低置信度知识可处理"),
    ], [1900, 4300, 3160], font_size=8.6)
    doc.add_heading("9.1 交互与性能策略", level=2)
    add_bullets(doc, [
        "首次只展示中心实体及一跳邻居，用户点击后继续展开；默认节点上限 150–300。",
        "节点按实体类型着色、关系按谓词分组，提供图例；悬停显示摘要，点击打开详情侧栏。",
        "查询结果保存视图状态，支持复制实体链接和导出当前子图 PNG/JSON。",
        "布局优先使用力导向与同心圆；大规模统计使用聚合图，不直接渲染全量节点。",
    ])

    add_page_break(doc)
    doc.add_heading("10. 安全、合规与可观测性", level=1)
    add_table(doc, ["风险面", "控制措施"], [
        ("数据合规", "建立来源白名单、许可记录、robots 检查、限速与删除机制。"),
        ("隐私", "禁止采集病例身份信息；入库前执行手机号、证件号等敏感信息检测。"),
        ("提示注入", "网页文本仅作为数据，不允许覆盖系统指令；工具调用使用白名单参数。"),
        ("模型幻觉", "必须返回原文证据；无证据事实不得入库；低置信度转人工。"),
        ("凭据安全", "API Key 通过环境变量/密钥文件注入，禁止提交到 Git。"),
        ("审计", "记录每条知识的来源、抽取器、模型版本、审核动作和写入时间。"),
    ], [2100, 7260], font_size=9)
    doc.add_heading("10.1 关键观测指标", level=2)
    add_bullets(doc, [
        "采集：请求成功率、平均延迟、空页率、重复率、各类别完成度。",
        "抽取：每千字实体数、三元组数、JSON 解析失败率、模型调用耗时与成本。",
        "图谱：新增/合并实体数、冲突数、孤立节点率、写入吞吐与失败数。",
        "应用：查询 P95 延迟、子图节点数、页面错误率、人工审核积压。",
    ])

    doc.add_heading("11. 实施计划与团队分工", level=1)
    add_table(doc, ["阶段", "周次", "主要任务", "里程碑"], [
        ("需求与本体", "第 1 周", "明确类别、来源、Schema、验收口径", "设计评审通过"),
        ("采集与清洗", "第 2 周", "实现来源适配器、去重、质量评分", "≥1,000 条样本"),
        ("抽取与融合", "第 3–4 周", "NER/RE/LLM、实体链接、评测集", "三元组可稳定产出"),
        ("图谱与 API", "第 5 周", "Neo4j 入库、查询、导出", "图查询接口通过"),
        ("前端展示", "第 6 周", "浏览、搜索、详情、路径、任务页", "完成端到端演示"),
        ("规模化与优化", "第 7 周", "跑满 120 类/3,600 条、修复质量问题", "数据指标达标"),
        ("交付与答辩", "第 8 周", "测试、README、PPT、演示视频", "完整交付包"),
    ], [1300, 1200, 4400, 2460], font_size=8.7)
    doc.add_heading("11.1 建议分工（4 人）", level=2)
    add_table(doc, ["角色", "主责", "协作事项"], [
        ("A：架构/后端", "编排、API、任务状态、部署", "统一 Schema 与集成"),
        ("B：数据/NLP", "采集、清洗、NER/RE、评测", "本体和提示词"),
        ("C：图谱", "Neo4j、实体融合、查询、导出", "质量规则与统计"),
        ("D：前端/展示", "可视化、搜索、详情、PPT", "端到端测试与演示"),
    ], [1800, 4100, 3460], font_size=8.8)

    doc.add_heading("12. 测试与验收方案", level=1)
    doc.add_heading("12.1 功能验收清单", level=2)
    add_bullets(doc, [
        "能够选择类别和来源启动构建任务，并看到各阶段状态与失败原因。",
        "清洗后有效类别数 ≥100、有效文本数 ≥3,000；建议交付值达到 120/3,600。",
        "能够从文本自动得到实体、属性和实体关系，且结果包含证据与置信度。",
        "能够把知识幂等写入图数据库，并导出节点、边、三元组和源文档数据文件。",
        "可视化页面支持图谱浏览、实体查询、实体详情和关系展示/路径查询。",
        "评测报告包含实体、关系、链接、可追溯率、重复率和流程成功率。",
        "换一小批新数据可进行增量构建，证明系统并非一次性静态演示。",
    ])
    doc.add_heading("12.2 非功能验收", level=2)
    add_table(doc, ["类别", "建议目标"], [
        ("查询性能", "常用实体搜索 P95 < 1 秒；一跳邻居查询 P95 < 2 秒。"),
        ("稳定性", "单批任务失败可重试；服务重启后可从检查点继续。"),
        ("可部署性", "Docker Compose 一键启动，README 从空环境可复现。"),
        ("可维护性", "采集器、模型与本体配置解耦；关键模块有单元测试。"),
        ("可解释性", "≥95% 入库关系可查看来源文档与证据片段。"),
    ], [2200, 7160], font_size=9)

    doc.add_heading("13. 项目交付物与目录建议", level=1)
    add_table(doc, ["交付物", "建议内容"], [
        ("完整源代码", "backend/、frontend/、agents/、extractors/、collectors/、tests/"),
        ("图谱数据", "categories.csv、documents.jsonl、entities.csv、relations.csv、triples.csv"),
        ("配置与部署", "configs/、ontology/、docker-compose.yml、.env.example"),
        ("文档", "README、系统设计、API、数据字典、评测与测试报告"),
        ("展示材料", "项目展示 PPT、演示脚本、可选演示视频"),
    ], [2400, 6960], font_size=9)
    doc.add_heading("13.1 PPT 建议结构（12–15 页）", level=2)
    add_numbers(doc, [
        "题目与项目目标；2. 痛点与方案亮点；3. 验收指标；4. 总体架构；5. 多智能体流程；",
        "6. 数据来源与规模；7. 本体设计；8. 混合抽取方案；9. 知识融合与质量控制；",
        "10. 图数据库与查询；11. 可视化界面；12. 评测结果；13. 现场演示流程；14. 总结与展望。",
    ])

    doc.add_heading("14. 主要风险与应对", level=1)
    add_table(doc, ["风险", "影响", "应对策略"], [
        ("100 类定义不清", "规模被质疑", "立项时固定类别表和统计脚本，类别可追溯到文档。"),
        ("来源不可抓取", "文本量不足", "准备多来源白名单和开放数据集兜底。"),
        ("LLM 成本或不可用", "抽取中断", "批处理、缓存、本地模型与规则降级。"),
        ("医疗知识幻觉", "知识错误", "证据强约束、置信度门槛、分层抽检。"),
        ("实体重复严重", "图谱碎片化", "别名词典、标准编码、向量召回 + 人工确认。"),
        ("大图渲染卡顿", "演示失败", "按需展开、节点上限、服务端聚合。"),
        ("最后才跑全量", "临近答辩暴露问题", "第 2 周开始持续跑数，每周冻结一版数据。"),
    ], [2000, 2200, 5160], font_size=8.7)

    doc.add_heading("15. 推荐的最小可行版本（MVP）", level=1)
    add_callout(doc, "MVP 边界", "先用 10 个类别、300 条文本贯通全链路，再横向扩展到 120 类、3,600 条。MVP 必须包含真实的自动采集/导入、结构化抽取、Neo4j 入库、实体查询和一跳关系展示，不以手工造图代替。")
    add_numbers(doc, [
        "冻结 ontology.yaml、document.schema.json 和 triple.schema.json。",
        "实现 1–2 个稳定数据源与本地文件导入器，完成去重和质量统计。",
        "实现规则 + 一种模型/LLM 的混合抽取，并保存证据与置信度。",
        "实现实体规范化、三元组校验和 Neo4j 幂等写入。",
        "完成查询 API 与图谱页面；最后补任务中心、审核页和规模化运行。",
    ])
    add_para(doc, "结论：该方案在题目要求之上增加了可追溯、可复核、可增量和可降级能力，既能突出“智能体驱动”的技术特色，也能用明确的数据文件、指标和页面完成客观验收。")

    # Ensure table rows do not split across pages.
    for table in doc.tables:
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            cant_split = OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)

    doc.core_properties.title = "智能体驱动的知识图谱自动构建系统设计方案"
    doc.core_properties.subject = "MedGraph 项目总体方案"
    doc.core_properties.author = "MedGraph 项目组"
    doc.core_properties.keywords = "知识图谱, 智能体, 信息抽取, 医疗文本, Neo4j"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
